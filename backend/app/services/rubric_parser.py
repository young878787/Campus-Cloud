"""Rubric parser service - Parse DOCX and PDF documents to Markdown."""

from __future__ import annotations

import io
from pathlib import Path


def parse_docx(file_bytes: bytes) -> str:
    """用 python-docx 解析 .docx，段落與表格都轉為 Markdown 文字。"""
    from docx import Document  # type: ignore

    doc = Document(io.BytesIO(file_bytes))
    lines: list[str] = []

    for block in _iter_block_items(doc):
        if block["type"] == "paragraph":
            text = block["text"].strip()
            if text:
                lines.append(text)
        elif block["type"] == "table":
            table = block["table"]
            # 收集所有格子文字
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append(cells)
            if not rows:
                continue
            # Markdown 表格
            header = rows[0]
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join(["---"] * len(header)) + " |")
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def parse_pdf(file_bytes: bytes) -> str:
    """用 pdfplumber 解析純文字型 PDF，表格轉為 Markdown，其餘為純文字（排除表格區域）。"""
    import pdfplumber  # type: ignore

    lines: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            # 先偵測此頁的表格與其 bbox，避免重複提取純文字
            tables = page.find_tables()
            table_bboxes: list[tuple[float, float, float, float]] = []
            for table in tables:
                data = table.extract()
                if not data:
                    continue
                table_bboxes.append(table.bbox)
                header = [str(cell or "").strip() for cell in data[0]]
                lines.append("| " + " | ".join(header) + " |")
                lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                for row in data[1:]:
                    cells = [str(cell or "").strip().replace("\n", " ") for cell in row]
                    lines.append("| " + " | ".join(cells) + " |")

            # 純文字（排除表格區域）：只保留不在任何表格 bbox 內的文字
            words = page.extract_words() or []

            def _in_any_table_bbox(word: dict) -> bool:
                x0 = float(word.get("x0", 0.0))
                x1 = float(word.get("x1", 0.0))
                top = float(word.get("top", 0.0))
                bottom = float(word.get("bottom", 0.0))
                for tb_x0, tb_top, tb_x1, tb_bottom in table_bboxes:
                    if (
                        x0 >= tb_x0
                        and x1 <= tb_x1
                        and top >= tb_top
                        and bottom <= tb_bottom
                    ):
                        return True
                return False

            non_table_words = [w for w in words if not _in_any_table_bbox(w)]

            # 依行號或 y 座標重建純文字行
            lines_by_key: dict[int, list[dict]] = {}
            for w in non_table_words:
                key = w.get("line_number")
                if key is None:
                    # 退而求其次：用 top 四捨五入當作行的 key
                    key = int(round(float(w.get("top", 0.0))))
                key = int(key)
                lines_by_key.setdefault(key, []).append(w)

            page_text_lines: list[str] = []
            for key in sorted(lines_by_key.keys()):
                words_in_line = sorted(
                    lines_by_key[key], key=lambda ww: float(ww.get("x0", 0.0))
                )
                text_line = " ".join(
                    str(ww.get("text", "")).strip()
                    for ww in words_in_line
                    if str(ww.get("text", "")).strip()
                )
                if text_line:
                    page_text_lines.append(text_line)

            page_text = "\n".join(page_text_lines)
            if page_text.strip():
                lines.append(page_text.strip())
    return "\n".join(lines)


def parse_document(filename: str, file_bytes: bytes) -> str:
    """根據副檔名選擇解析器，回傳純文字（Markdown 格式）。"""
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        return parse_docx(file_bytes)
    elif suffix == ".pdf":
        return parse_pdf(file_bytes)
    else:
        raise ValueError(f"不支援的文件格式：{suffix}（目前支援 .docx / .pdf）")


# ──────────────────────────────────────────────────────
# python-docx 工具：按文件順序迭代段落 + 表格
# ──────────────────────────────────────────────────────


def _iter_block_items(doc):  # type: ignore
    """
    按文件原始順序，逐一 yield 段落或表格。
    python-docx 的 doc.paragraphs 和 doc.tables 是分開的，
    此函式利用 XML element 順序確保正確排列。
    """
    from docx.oxml.ns import qn  # type: ignore

    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            from docx.text.paragraph import Paragraph  # type: ignore

            yield {"type": "paragraph", "text": Paragraph(child, doc).text}
        elif tag == "tbl":
            from docx.table import Table  # type: ignore

            yield {"type": "table", "table": Table(child, doc)}
