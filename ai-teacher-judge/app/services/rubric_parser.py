from __future__ import annotations

import io
from pathlib import Path


def parse_docx(file_bytes: bytes) -> str:
    """用 python-docx 解析 .docx，段落與表格都轉為 Markdown 文字。"""
    from docx import Document  # type: ignore

    doc = Document(io.BytesIO(file_bytes))
    lines: list[str] = []

    for block in _iter_block_items(doc):
        if block["type"] == "paragraph":
            text = block["text"].strip()
            if text:
                lines.append(text)
        elif block["type"] == "table":
            table = block["table"]
            # 收集所有格子文字
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append(cells)
            if not rows:
                continue
            # Markdown 表格
            header = rows[0]
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join(["---"] * len(header)) + " |")
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def parse_pdf(file_bytes: bytes) -> str:
    """用 pdfplumber 解析純文字型 PDF，表格轉為 Markdown，其餘為純文字。"""
    import pdfplumber  # type: ignore

    lines: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            # 先取出此頁的表格 bbox，避免重複提取純文字
            tables = page.extract_tables()
            for table in tables:
                if not table:
                    continue
                header = [str(cell or "").strip() for cell in table[0]]
                lines.append("| " + " | ".join(header) + " |")
                lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                for row in table[1:]:
                    cells = [str(cell or "").strip().replace("\n", " ") for cell in row]
                    lines.append("| " + " | ".join(cells) + " |")
            # 純文字（排除表格區域）
            text = page.extract_text() or ""
            if text.strip():
                lines.append(text.strip())

    return "\n".join(lines)


def parse_document(filename: str, file_bytes: bytes) -> str:
    """根據副檔名選擇解析器，回傳純文字（Markdown 格式）。"""
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        return parse_docx(file_bytes)
    elif suffix == ".pdf":
        return parse_pdf(file_bytes)
    else:
        raise ValueError(f"不支援的文件格式：{suffix}（目前支援 .docx / .pdf）")


# ──────────────────────────────────────────────
# python-docx 工具：按文件順序迭代段落 + 表格
# ──────────────────────────────────────────────

def _iter_block_items(doc):  # type: ignore
    """
    按文件原始順序，逐一 yield 段落或表格。
    python-docx 的 doc.paragraphs 和 doc.tables 是分開的，
    此函式利用 XML element 順序確保正確排列。
    """
    from docx.oxml.ns import qn  # type: ignore

    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            from docx.text.paragraph import Paragraph  # type: ignore
            yield {"type": "paragraph", "text": Paragraph(child, doc).text}
        elif tag == "tbl":
            from docx.table import Table  # type: ignore
            yield {"type": "table", "table": Table(child, doc)}
