"""
文件處理工具 - 提取文件內容並轉換為結構化格式
支援 DOCX, PDF, TXT 等格式
圖片以占位符形式標記
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Union, BinaryIO

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[Warning] python-docx 未安裝，.docx 支援將不可用")
    print("[Info] 安裝: pip install python-docx")

try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("[Warning] pypdf 未安裝，.pdf 支援將不可用")
    print("[Info] 安裝: pip install pypdf")


class DocumentProcessor:
    """
    文件處理器
    支援多種文件格式的內容提取和結構化
    """

    @staticmethod
    def detect_file_type(file_path: Union[str, Path, None] = None, filename: str = None) -> str:
        """
        檢測文件類型
        
        Args:
            file_path: 文件路徑
            filename: 文件名稱
            
        Returns:
            文件類型 ("docx", "pdf", "txt", "unknown")
        """
        name = filename if filename else str(file_path)
        name_lower = name.lower()
        
        if name_lower.endswith('.docx'):
            return 'docx'
        elif name_lower.endswith('.doc'):
            return 'doc'  # 舊格式，需要轉換
        elif name_lower.endswith('.pdf'):
            return 'pdf'
        elif name_lower.endswith('.txt'):
            return 'txt'
        elif name_lower.endswith('.md'):
            return 'markdown'
        else:
            return 'unknown'

    @staticmethod
    def extract_text_from_docx(file_data: Union[bytes, BinaryIO, str, Path]) -> str:
        """
        從 DOCX 文件提取文字和結構
        圖片以占位符形式標記
        
        Args:
            file_data: 文件數據（bytes、文件對象或路徑）
            
        Returns:
            Markdown 格式的文件內容
            
        Raises:
            ImportError: python-docx 未安裝
            Exception: 文件處理錯誤
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx 未安裝，無法處理 .docx 文件\n"
                "請執行: pip install python-docx"
            )
        
        try:
            # 處理不同輸入類型
            if isinstance(file_data, (str, Path)):
                doc = Document(file_data)
            elif isinstance(file_data, bytes):
                doc = Document(io.BytesIO(file_data))
            else:
                doc = Document(file_data)
            
            content_parts = []
            image_counter = 0
            
            # 提取段落
            for para in doc.paragraphs:
                text = para.text.strip()
                
                if not text:
                    # 檢查是否包含圖片
                    if any(run._element.xpath('.//a:blip') for run in para.runs):
                        image_counter += 1
                        content_parts.append(f"\n📷 **[圖片 {image_counter}]**\n")
                    continue
                
                # 根據樣式判斷標題層級
                style_name = para.style.name.lower() if para.style else ""
                
                if 'heading 1' in style_name or 'title' in style_name:
                    content_parts.append(f"\n# {text}\n")
                elif 'heading 2' in style_name:
                    content_parts.append(f"\n## {text}\n")
                elif 'heading 3' in style_name:
                    content_parts.append(f"\n### {text}\n")
                elif 'heading' in style_name:
                    content_parts.append(f"\n#### {text}\n")
                else:
                    # 普通段落
                    content_parts.append(f"{text}\n")
            
            # 提取表格
            for table_idx, table in enumerate(doc.tables, 1):
                content_parts.append(f"\n**[表格 {table_idx}]**\n")
                
                # 轉換為 Markdown 表格
                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    content_parts.append("| " + " | ".join(cells) + " |")
                    
                    # 表頭分隔線
                    if row_idx == 0:
                        content_parts.append("| " + " | ".join(['---'] * len(cells)) + " |")
                
                content_parts.append("\n")
            
            result = '\n'.join(content_parts).strip()
            
            # 添加摘要信息
            summary = f"*[文件包含 {len(doc.paragraphs)} 個段落、{len(doc.tables)} 個表格、{image_counter} 張圖片]*\n\n"
            
            return summary + result
        
        except Exception as e:
            raise Exception(f"DOCX 文件處理失敗: {str(e)}")

    @staticmethod
    def extract_text_from_pdf(file_data: Union[bytes, BinaryIO, str, Path]) -> str:
        """
        從 PDF 文件提取文字
        圖片以占位符形式標記
        
        Args:
            file_data: 文件數據（bytes、文件對象或路徑）
            
        Returns:
            Markdown 格式的文件內容
            
        Raises:
            ImportError: pypdf 未安裝
            Exception: 文件處理錯誤
        """
        if not PDF_AVAILABLE:
            raise ImportError(
                "pypdf 未安裝，無法處理 .pdf 文件\n"
                "請執行: pip install pypdf"
            )
        
        try:
            # 處理不同輸入類型
            if isinstance(file_data, (str, Path)):
                reader = PdfReader(file_data)
            elif isinstance(file_data, bytes):
                reader = PdfReader(io.BytesIO(file_data))
            else:
                reader = PdfReader(file_data)
            
            content_parts = []
            total_images = 0
            
            # 逐頁提取
            for page_num, page in enumerate(reader.pages, 1):
                # 提取文字
                text = page.extract_text()
                
                if text.strip():
                    content_parts.append(f"\n---\n**[第 {page_num} 頁]**\n\n{text}\n")
                
                # 統計圖片數量（安全存取，避免缺少 Resources 時 KeyError）
                resources = page.get('/Resources', {})
                if '/XObject' in resources:
                    xobject = resources['/XObject'].get_object()
                    images_in_page = sum(
                        1 for obj in xobject
                        if xobject[obj].get('/Subtype') == '/Image'
                    )
                    if images_in_page > 0:
                        content_parts.append(f"\n📷 *[此頁包含 {images_in_page} 張圖片]*\n")
                        total_images += images_in_page
            
            result = '\n'.join(content_parts).strip()
            
            # 添加摘要信息
            metadata = reader.metadata
            summary_parts = [f"*[PDF 文件：{len(reader.pages)} 頁"]
            
            if metadata:
                if metadata.title:
                    summary_parts.append(f"標題：{metadata.title}")
                if metadata.author:
                    summary_parts.append(f"作者：{metadata.author}")
            
            if total_images > 0:
                summary_parts.append(f"共 {total_images} 張圖片")
            
            summary = ", ".join(summary_parts) + "]*\n\n"
            
            return summary + result
        
        except Exception as e:
            raise Exception(f"PDF 文件處理失敗: {str(e)}")

    @staticmethod
    def extract_text_from_txt(file_data: Union[bytes, BinaryIO, str, Path]) -> str:
        """
        從 TXT 文件提取文字
        
        Args:
            file_data: 文件數據（bytes、文件對象或路徑）
            
        Returns:
            文件內容
        """
        try:
            if isinstance(file_data, (str, Path)):
                with open(file_data, 'r', encoding='utf-8') as f:
                    return f.read()
            elif isinstance(file_data, bytes):
                # 嘗試多種編碼
                for encoding in ['utf-8', 'gbk', 'big5', 'latin-1']:
                    try:
                        return file_data.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                # 如果所有編碼都失敗，使用 utf-8 並忽略錯誤
                return file_data.decode('utf-8', errors='ignore')
            else:
                content = file_data.read()
                if isinstance(content, bytes):
                    return content.decode('utf-8', errors='ignore')
                return content
        
        except Exception as e:
            raise Exception(f"TXT 文件處理失敗: {str(e)}")

    @classmethod
    def extract_document(
        cls,
        file_data: Union[bytes, BinaryIO, str, Path],
        filename: str = None
    ) -> dict:
        """
        智能提取文件內容
        根據文件類型自動選擇處理方法
        
        Args:
            file_data: 文件數據
            filename: 文件名（用於類型檢測）
            
        Returns:
            {
                "content": "文件內容（Markdown 格式）",
                "file_type": "文件類型",
                "success": True/False,
                "error": "錯誤信息（如果失敗）"
            }
        """
        try:
            # 檢測文件類型
            file_type = cls.detect_file_type(filename=filename)
            
            # 根據類型提取內容
            if file_type == 'docx':
                content = cls.extract_text_from_docx(file_data)
            elif file_type == 'pdf':
                content = cls.extract_text_from_pdf(file_data)
            elif file_type in ['txt', 'markdown']:
                content = cls.extract_text_from_txt(file_data)
            elif file_type == 'doc':
                return {
                    "content": "",
                    "file_type": file_type,
                    "success": False,
                    "error": "不支援 .doc 格式，請轉換為 .docx"
                }
            else:
                return {
                    "content": "",
                    "file_type": file_type,
                    "success": False,
                    "error": f"不支援的文件格式: {filename}"
                }
            
            return {
                "content": content,
                "file_type": file_type,
                "success": True,
                "error": None
            }
        
        except ImportError as e:
            return {
                "content": "",
                "file_type": file_type if 'file_type' in locals() else 'unknown',
                "success": False,
                "error": str(e)
            }
        except Exception as e:
            return {
                "content": "",
                "file_type": file_type if 'file_type' in locals() else 'unknown',
                "success": False,
                "error": f"文件處理失敗: {str(e)}"
            }


def create_document_prompt(
    document_content: str,
    user_message: str,
    file_type: str = "文件"
) -> str:
    """
    創建文件分析的 Prompt
    使用結構化格式區分文件內容和用戶問題
    
    Args:
        document_content: 提取的文件內容
        user_message: 用戶問題
        file_type: 文件類型描述
        
    Returns:
        格式化的 Prompt
    """
    prompt = f"""# 📄 {file_type.upper()} 文件內容

{document_content}

---

# ❓ 用戶問題

{user_message}

---

請基於上述文件內容回答用戶問題。如果文件中沒有相關信息，請明確說明。"""
    
    return prompt


# 便捷函數
def extract_document(file_data, filename: str = None) -> dict:
    """
    便捷函數：提取文件內容
    
    Args:
        file_data: 文件數據
        filename: 文件名
        
    Returns:
        提取結果字典
    """
    return DocumentProcessor.extract_document(file_data, filename)
